import re
import docx
from docx.oxml.ns import qn


def simplify_layout_tables(doc: docx.Document) -> docx.Document:
    """
    Converts sparse layout tables (commonly used on title pages) into plain
    paragraphs before Pandoc conversion. This prevents malformed wide pipe tables
    with mostly empty cells in Markdown output.
    """
    for table in list(doc.tables):
        if not _is_layout_table(table):
            continue

        tbl = table._tbl
        extracted_lines = []
        for row in table.rows:
            row_texts = []
            seen_cells = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)

                text = " ".join(part.strip() for part in cell.text.splitlines() if part.strip())
                if text:
                    row_texts.append(text)

            if row_texts:
                extracted_lines.append(" — ".join(row_texts))

        for line in extracted_lines:
            _insert_paragraph_before(tbl, line)

        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)

    return doc


def _is_layout_table(table) -> bool:
    total_cells = 0
    non_empty_cells = 0
    image_cells = 0

    for row in table.rows:
        seen_cells = set()
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)

            total_cells += 1
            text = " ".join(part.strip() for part in cell.text.splitlines() if part.strip())
            if text:
                non_empty_cells += 1
            if cell._tc.xpath('.//w:drawing | .//w:pict'):
                image_cells += 1

    if total_cells == 0:
        return False

    non_empty_ratio = non_empty_cells / total_cells
    return non_empty_ratio <= 0.45 and (image_cells > 0 or non_empty_cells <= 3)


def _insert_paragraph_before(anchor, text: str) -> None:
    new_p = docx.oxml.OxmlElement('w:p')
    new_r = docx.oxml.OxmlElement('w:r')
    new_t = docx.oxml.OxmlElement('w:t')
    if text != text.strip():
        new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    anchor.addprevious(new_p)

def remove_toc_and_add_placeholder(doc: docx.Document) -> docx.Document:
    """
    Finds standard (w:sdt) and manual (md2gost) TOC in the docx file
    and replaces them with [TOC]. Ensures both variants are supported safely.
    """
    toc_inserted = False
    
    for sdt in doc._body._element.xpath('//w:sdt'):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is not None:
            is_toc = False
            for child in sdtPr:
                if child.tag == qn('w:docPartObj'):
                    for sub in child:
                        if sub.tag == qn('w:docPartGallery') and sub.get(qn('w:val')) == 'Table of Contents':
                            is_toc = True
                            break
            
            if is_toc:
                if not toc_inserted:
                    new_p = docx.oxml.OxmlElement('w:p')
                    new_r = docx.oxml.OxmlElement('w:r')
                    new_t = docx.oxml.OxmlElement('w:t')
                    new_t.text = '[TOC]'
                    new_r.append(new_t)
                    new_p.append(new_r)
                    sdt.addprevious(new_p)
                    toc_inserted = True
                
                sdt.getparent().remove(sdt)

    toc_pattern = re.compile(r'^.+\t\d+$')
    
    for p in doc.paragraphs:
        is_toc_style = p.style.name.startswith('TOC') or 'toc' in (p.style.name or '').lower()
        is_manual_toc = bool(toc_pattern.match(p.text.strip()))
        has_hyperlink = len(p._element.xpath('.//w:hyperlink')) > 0
        
        if is_toc_style or (is_manual_toc and has_hyperlink):
            if not toc_inserted:
                p.text = '[TOC]'
                try:
                    p.style = doc.styles['Normal']
                except KeyError:
                    pass
                toc_inserted = True
            else:
                p_ele = p._element
                if p_ele.getparent() is not None:
                    p_ele.getparent().remove(p_ele)

    return doc

def prepare_tables_for_markdown(doc: docx.Document) -> docx.Document:
    """
    Flattens multiline cells into single lines and removes col/row spans
    so that pypandoc can safely generate standard markdown pipe tables
    instead of [TABLE] placeholders.
    """
    for table in doc.tables:
        for row in table.rows:
            seen_cells = set()
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                
                # Remove colspans so pandoc doesn't trip [TABLE]
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    tcPr.remove(gridSpan)
                
                # Remove rowspans
                vMerge = tcPr.find(qn('w:vMerge'))
                if vMerge is not None:
                    tcPr.remove(vMerge)

                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                
                # Flatten paragraphs
                if len(cell.paragraphs) > 1:
                    first_p = cell.paragraphs[0]._p
                    for p in cell.paragraphs[1:]:
                        space_run = docx.oxml.OxmlElement('w:r')
                        space_text = docx.oxml.OxmlElement('w:t')
                        space_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        space_text.text = " <br> "
                        space_run.append(space_text)
                        
                        first_p.append(space_run)
                        
                        for child in list(p._p):
                            if child.tag != qn("w:pPr"):
                                first_p.append(child)
                        p._p.getparent().remove(p._p)
    return doc
